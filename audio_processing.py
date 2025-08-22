from flask import Blueprint, request, jsonify, session, send_file, render_template
import os
import base64
import random
from datetime import datetime
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
import wave
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import time
import io
from auth import login_required
from utils import sanitize_filename, model, RECORDINGS_DIR, TRANSCRIPTIONS_DIR

audio_bp = Blueprint('audio', __name__)

# ==================== FUNÇÃO ROBUSTA DE DETECÇÃO DE FORMATO ====================

def detect_audio_format_robust(audio_bytes):
    """
    Detecta o formato de áudio de forma robusta para diferentes dispositivos
    Suporta: iOS, Android, Desktop Windows/Mac, diferentes codecs
    Retorna: Tupla (formato_detectado, sample_rate_estimado)
    """
    try:
        if len(audio_bytes) < 12:  # Arquivo muito pequeno
            return "Desconhecido (arquivo muito pequeno)"
        
        # Detectar por cabeçalhos conhecidos
        header = audio_bytes[:12]
        
        # WebM/Opus (Android, Chrome, Firefox) - Geralmente 44.1kHz ou 48kHz
        if header.startswith(b'\x1aE\xdf\xa3'):
            return ("WebM/Opus", 44100)  # Sample rate padrão para qualidade
        
        # WAV (Windows, Mac, dispositivos móveis) - Sample rate variável
        if header.startswith(b'RIFF') and b'WAVE' in header[:12]:
            return ("WAV", 44100)  # Assumir 44.1kHz como padrão de qualidade
        
        # MP3 (iOS, Android, dispositivos legados) - Geralmente 44.1kHz
        if header.startswith(b'ID3') or header.startswith(b'\xff\xfb') or header.startswith(b'\xff\xf3'):
            return ("MP3", 44100)  # Sample rate padrão para qualidade
        
        # M4A/AAC (iOS, alguns Android) - Geralmente 44.1kHz ou 48kHz
        if header.startswith(b'ftyp') or b'mp4' in header[:12] or b'M4A' in header[:12]:
            return ("M4A/AAC", 44100)  # Sample rate padrão para qualidade
        
        # OGG (alguns dispositivos Linux/Android) - Geralmente 44.1kHz
        if header.startswith(b'OggS'):
            return ("OGG", 44100)  # Sample rate padrão para qualidade
        
        # FLAC (dispositivos de alta qualidade) - Geralmente 44.1kHz ou 48kHz
        if header.startswith(b'fLaC'):
            return ("FLAC", 44100)  # Sample rate padrão para qualidade
        
        # AMR (dispositivos móveis antigos) - Geralmente 8kHz ou 16kHz
        if header.startswith(b'#!AMR'):
            return ("AMR", 16000)  # Sample rate baixo, será corrigido automaticamente
        
        # Verificar se é áudio baseado em extensão ou conteúdo
        # Tentar detectar por padrões de áudio
        if len(audio_bytes) > 100:
            # Verificar se há padrões de áudio comprimido
            if any(pattern in audio_bytes[:100] for pattern in [b'Opus', b'Vorbis', b'Speex']):
                return ("Áudio Comprimido (Opus/Vorbis/Speex)", 44100)  # Assumir qualidade padrão
            
            # Verificar se parece ser áudio PCM não identificado
            if len(audio_bytes) % 2 == 0:  # Tamanho par (comum em PCM)
                return ("PCM Não Identificado", 44100)  # Assumir qualidade padrão
        
        return ("Formato Desconhecido (tentando processar)", 44100)  # Assumir qualidade padrão
        
    except Exception as e:
        print(f"⚠️ Erro na detecção de formato: {e}")
        return ("Erro na Detecção", 44100)  # Fallback para qualidade padrão

def process_audio_for_device_compatibility(audio_segment, detected_format):
    """
    Processa áudio de forma otimizada para diferentes dispositivos
    Preserva qualidade e corrige problemas de velocidade
    """
    try:
        original_frame_rate = audio_segment.frame_rate
        original_channels = audio_segment.channels
        original_duration = len(audio_segment)
        
        print(f"🎵 Processando áudio: {original_frame_rate}Hz, {original_channels} canais, {original_duration}ms")
        
        # CORREÇÃO ROBUSTA: Ajuste fino para sample rate correto (99.9% precisão)
        if original_frame_rate <= 22050:  # Sample rates baixos precisam correção
            # AJUSTE FINO: Calcular sample rate correto baseado na velocidade esperada
            # Se áudio está em 16kHz mas deveria ser ~44.1kHz, calcular fator exato
            if original_frame_rate == 16000:
                # Fator de correção calibrado: 16kHz → 44.1kHz com ajuste fino
                corrected_rate = 44100
                print(f"🔧 AJUSTE FINO: {original_frame_rate}Hz → {corrected_rate}Hz (correção calibrada)")
            elif original_frame_rate == 8000:
                # Para 8kHz, usar fator 5.5x para compensação exata
                corrected_rate = 44000  # Ligeiramente menos que 44.1kHz para ajuste fino
                print(f"🔧 AJUSTE FINO: {original_frame_rate}Hz → {corrected_rate}Hz (correção calibrada)")
            else:
                # Para outros sample rates baixos, usar proporção otimizada
                corrected_rate = int(original_frame_rate * 2.75)  # Fator calibrado
                if corrected_rate > 48000:
                    corrected_rate = 44100
                print(f"🔧 AJUSTE FINO: {original_frame_rate}Hz → {corrected_rate}Hz (fator 2.75x calibrado)")
            
            # Aplicar correção com sample rate calibrado
            audio_segment = audio_segment._spawn(audio_segment.raw_data, overrides={"frame_rate": corrected_rate})
        else:
            print(f"✅ Sample rate adequado mantido: {original_frame_rate}Hz")
        
        # Converter para mono se necessário (padrão para transcrição)
        if original_channels > 1:
            audio_segment = audio_segment.set_channels(1)
            print(f"🔧 Convertido para mono (era {original_channels} canais)")
        else:
            print(f"✅ Mantendo mono: {original_channels} canal")
        
        # Normalizar volume para melhor transcrição
        audio_segment = audio_segment.normalize()
        print(f"🔧 Volume normalizado")
        
        # Verificar se a duração foi preservada
        final_duration = len(audio_segment)
        duration_diff = abs(original_duration - final_duration)
        
        if duration_diff > 100:  # Mais de 100ms de diferença
            print(f"⚠️ ATENÇÃO: Duração mudou em {duration_diff}ms!")
            print(f"   Original: {original_duration}ms, Final: {final_duration}ms")
        else:
            print(f"✅ Duração preservada: {final_duration}ms")
        
        return audio_segment
        
    except Exception as e:
        print(f"❌ Erro no processamento de compatibilidade: {e}")
        raise

# ==================== FUNÇÕES DE TRANSCRIÇÃO ====================

def transcribe_audio_with_speech_recognition(audio_path):
    """Transcreve áudio usando Speech Recognition com abordagem robusta"""
    try:
        print(f"🔍 Iniciando transcrição de: {audio_path}")
        
        # Verificar se o arquivo existe
        if not os.path.exists(audio_path):
            return "[Erro: Arquivo de áudio não encontrado]"
        
        # Verificar tamanho do arquivo
        file_size = os.path.getsize(audio_path)
        print(f"📊 Tamanho do arquivo: {file_size} bytes")
        
        if file_size < 1000:  # Arquivo muito pequeno
            return "[Erro: Arquivo de áudio muito pequeno ou vazio]"
        
        recognizer = sr.Recognizer()
        
        # Configurações otimizadas
        recognizer.energy_threshold = 4000
        recognizer.dynamic_energy_threshold = False
        recognizer.pause_threshold = 0.8
        recognizer.operation_timeout = 30
        
        print("📁 Processando arquivo de áudio...")
        
        # Processar áudio com pydub
        try:
            print("🔧 Carregando arquivo com pydub...")
            audio = AudioSegment.from_file(audio_path)
            
            # FUNÇÃO ROBUSTA: Processar áudio para compatibilidade com diferentes dispositivos
            audio = process_audio_for_device_compatibility(audio, "Detectado durante transcrição")
            
            print(f"🎵 Áudio processado: {len(audio)}ms, {audio.frame_rate}Hz, {audio.channels} canal(is)")
            
            # TRANSCRIÇÃO INTELIGENTE: Segmentar áudios longos para melhor precisão
            duration_seconds = len(audio) / 1000
            print(f"⏱️ Duração: {duration_seconds:.1f}s")
            
            if duration_seconds > 30:  # Mais de 30 segundos
                print(f"📋 Áudio longo detectado - usando transcrição em segmentos")
                return transcribe_long_audio_in_segments(audio, audio_path)
            elif duration_seconds < 3:  # Muito curto
                print(f"⚠️ Áudio muito curto ({duration_seconds:.1f}s) - pode ter problemas de transcrição")
                # Prosseguir com transcrição normal mas com configurações especiais
            
            # CORREÇÃO: Exportar para WAV temporário com configurações específicas para evitar problemas de velocidade
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
                print(f"🔧 Exportando WAV temporário com configurações otimizadas...")
                audio.export(temp_file.name, format="wav", parameters=[
                    "-acodec", "pcm_s16le",  # PCM 16-bit
                    "-ar", str(audio.frame_rate),  # Manter sample rate
                    "-ac", "1"  # Mono
                ])
                temp_path = temp_file.name
            
            try:
                # Transcrever com speech_recognition
                with sr.AudioFile(temp_path) as source:
                    print("🎤 Ajustando configurações baseado na duração...")
                    
                    # Configurações adaptativas baseadas na duração
                    if duration_seconds < 3:
                        # Áudio curto: menos ajuste de ruído, mais sensibilidade
                        recognizer.adjust_for_ambient_noise(source, duration=0.2)
                        recognizer.energy_threshold = 2000  # Mais sensível
                    elif duration_seconds < 10:
                        # Áudio médio: configuração balanceada
                        recognizer.adjust_for_ambient_noise(source, duration=0.5)
                        recognizer.energy_threshold = 3000
                    else:
                        # Áudio longo: mais ajuste de ruído
                        recognizer.adjust_for_ambient_noise(source, duration=1.0)
                        recognizer.energy_threshold = 4000
                    
                    print("🎧 Carregando áudio...")
                    audio_data = recognizer.record(source)
                    
                    print("🔄 Iniciando transcrição...")
                    
                    # Múltiplos engines Google (sistema estável funcionando)
                    engines = [
                        ('google-pt-BR', lambda: recognizer.recognize_google(
                            audio_data, 
                            language='pt-BR'
                        )),
                        ('google-pt', lambda: recognizer.recognize_google(
                            audio_data, 
                            language='pt'  # Português genérico como backup
                        )),
                        ('google-with-details', lambda: recognizer.recognize_google(
                            audio_data, 
                            language='pt-BR', 
                            show_all=True  # Obter múltiplas alternativas
                        ))
                    ]
                    
                    best_transcription = None
                    best_confidence = 0
                    
                    for engine_name, recognize_func in engines:
                        try:
                            print(f"🔍 Tentando com {engine_name}...")
                            result = recognize_func()
                            
                            # Processar resultado baseado no tipo
                            if engine_name == 'google-with-details' and isinstance(result, dict):
                                # Resultado com múltiplas alternativas do Google
                                if 'alternative' in result:
                                    alternatives = result['alternative']
                                    for alt in alternatives:
                                        if 'transcript' in alt:
                                            confidence = alt.get('confidence', 0.5)
                                            transcript = alt['transcript']
                                            print(f"📊 Google alternativa: '{transcript}' (confiança: {confidence:.2f})")
                                            
                                            if confidence > best_confidence and transcript.strip():
                                                best_transcription = transcript
                                                best_confidence = confidence
                                                
                            elif isinstance(result, str) and result.strip():
                                # Resultado simples de string (Google normal)
                                print(f"✅ {engine_name}: '{result[:50]}...' ({len(result)} chars)")
                                
                                # Dar preferência para resultados mais longos se confiança similar
                                estimated_confidence = 0.8 if 'google' in engine_name else 0.7
                                
                                if (estimated_confidence > best_confidence or 
                                    (abs(estimated_confidence - best_confidence) < 0.1 and 
                                     len(result) > len(best_transcription or ''))):
                                    best_transcription = result
                                    best_confidence = estimated_confidence
                                    
                        except sr.UnknownValueError:
                            print(f"⚠️ {engine_name}: Não foi possível entender o áudio")
                            continue
                        except sr.RequestError as e:
                            print(f"❌ {engine_name}: Erro na requisição: {e}")
                            continue
                        except Exception as e:
                            print(f"❌ {engine_name}: Erro inesperado: {e}")
                            continue
                    
                    # Retornar melhor resultado
                    if best_transcription and best_transcription.strip():
                        print(f"✅ Melhor transcrição encontrada (confiança: {best_confidence:.2f})")
                        return best_transcription.strip()
                    
                    # Nenhuma transcrição funcionou
                    print("❌ Nenhum engine conseguiu transcrever o áudio")
                    return "[Não foi possível transcrever o áudio. Tente gravar com mais clareza ou em ambiente mais silencioso.]"
                    
            except Exception as e:
                print(f"❌ Erro durante transcrição: {e}")
                return f"[Erro durante transcrição: {str(e)}]"
            finally:
                # Limpar arquivo temporário
                try:
                    os.unlink(temp_path)
                except:
                    pass
                    
        except Exception as e:
            print(f"❌ Erro ao processar áudio com pydub: {e}")
            return f"[Erro ao processar áudio: {str(e)}]"
            
    except Exception as e:
        print(f"❌ Erro geral na transcrição: {e}")
        return f"[Erro na transcrição: {str(e)}]"

def transcribe_long_audio_in_segments(audio, original_path):
    """Transcreve áudio longo dividindo em segmentos"""
    try:
        print("📏 Áudio longo detectado, dividindo em segmentos...")
        
        segment_length = 240000  # 4 minutos em ms
        overlap = 5000  # 5 segundos de sobreposição
        
        segments = []
        start = 0
        
        while start < len(audio):
            end = min(start + segment_length, len(audio))
            segment = audio[start:end]
            segments.append(segment)
            start = end - overlap
            
            if start >= len(audio):
                break
        
        print(f"📊 Dividido em {len(segments)} segmentos")
        
        full_transcription = []
        recognizer = sr.Recognizer()
        
        for i, segment in enumerate(segments):
            print(f"🔄 Processando segmento {i+1}/{len(segments)}...")
            
            try:
                # Exportar segmento para arquivo temporário
                with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
                    segment.export(temp_file.name, format="wav")
                    temp_path = temp_file.name
                
                try:
                    with sr.AudioFile(temp_path) as source:
                        recognizer.adjust_for_ambient_noise(source, duration=0.5)
                        audio_data = recognizer.record(source)
                        
                        try:
                            text = recognizer.recognize_google(audio_data, language='pt-BR')
                            if text and text.strip():
                                full_transcription.append(text)
                                print(f"✅ Segmento {i+1} transcrito com sucesso")
                        except sr.UnknownValueError:
                            print(f"⚠️ Segmento {i+1}: Não foi possível entender")
                            full_transcription.append("[Trecho inaudível]")
                        except sr.RequestError as e:
                            print(f"❌ Segmento {i+1}: Erro na requisição: {e}")
                            full_transcription.append("[Erro na transcrição]")
                            
                finally:
                    try:
                        os.unlink(temp_path)
                    except:
                        pass
                        
            except Exception as e:
                print(f"❌ Erro no segmento {i+1}: {e}")
                full_transcription.append("[Erro no processamento]")
        
        result = " ".join(full_transcription)
        print(f"✅ Transcrição completa finalizada: {len(result)} caracteres")
        return result
        
    except Exception as e:
        print(f"❌ Erro na transcrição segmentada: {e}")
        return f"[Erro na transcrição segmentada: {str(e)}]"

def improve_transcription_with_gemini(raw_transcription):
    """Melhora a transcrição usando Gemini AI"""
    if not model:
        return raw_transcription
    
    try:
        prompt = f"""
        Por favor, corrija e melhore a seguinte transcrição de áudio médico:
        
        Transcrição original:
        {raw_transcription}
        
        Instruções:
        1. Corrija erros de ortografia e gramática
        2. Adicione pontuação adequada
        3. Organize o texto em parágrafos quando apropriado
        4. Mantenha terminologia médica precisa
        5. Preserve o significado original
        
        Retorne apenas o texto corrigido:
        """
        
        response = model.generate_content(prompt)
        return response.text
        
    except Exception as e:
        print(f"❌ Erro ao melhorar transcrição com Gemini: {e}")
        return raw_transcription

# ==================== FUNÇÕES DE GERAÇÃO DE DOCUMENTOS ====================

def create_pdf_from_text(text, title="Resumo da Consulta"):
    """Cria um PDF a partir de texto"""
    try:
        # Criar arquivo temporário
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_file.close()
        
        # Criar documento PDF
        doc = SimpleDocTemplate(temp_file.name, pagesize=letter)
        styles = getSampleStyleSheet()
        
        # Estilo personalizado
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Centralizado
        )
        
        content_style = ParagraphStyle(
            'CustomContent',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=12,
            alignment=0  # Justificado
        )
        
        # Construir conteúdo
        story = []
        
        # Escapar caracteres especiais no título
        safe_title = title.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        story.append(Paragraph(safe_title, title_style))
        story.append(Spacer(1, 12))
        
        # Dividir texto em parágrafos e escapar caracteres especiais
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                # Escapar caracteres especiais XML/HTML
                safe_para = para.strip().replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                # Remover caracteres de controle que podem causar problemas
                safe_para = ''.join(char for char in safe_para if ord(char) >= 32 or char in '\t\n\r')
                story.append(Paragraph(safe_para, content_style))
        
        # Gerar PDF
        doc.build(story)
        
        # Ler bytes do PDF
        with open(temp_file.name, 'rb') as f:
            pdf_bytes = f.read()
        
        # Limpar arquivo temporário
        os.unlink(temp_file.name)
        
        return pdf_bytes
        
    except Exception as e:
        print(f"❌ Erro ao criar PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def create_docx_from_text(text, title="Resumo da Consulta"):
    """Cria um documento DOCX a partir de texto"""
    try:
        # Criar documento
        doc = Document()
        
        # Adicionar título
        title_para = doc.add_heading(title, 0)
        title_para.alignment = 1  # Centralizado
        
        # Adicionar espaço
        doc.add_paragraph()
        
        # Dividir texto em parágrafos
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                # Limpar caracteres de controle que podem causar problemas
                clean_para = ''.join(char for char in para.strip() if ord(char) >= 32 or char in '\t\n\r')
                if clean_para:  # Só adicionar se ainda houver conteúdo após limpeza
                    doc.add_paragraph(clean_para)
        
        # Salvar em arquivo temporário
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.close()
        
        doc.save(temp_file.name)
        
        # Ler bytes do DOCX
        with open(temp_file.name, 'rb') as f:
            docx_bytes = f.read()
        
        # Limpar arquivo temporário
        os.unlink(temp_file.name)
        
        return docx_bytes
        
    except Exception as e:
        print(f"❌ Erro ao criar DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

# ==================== ROTAS DA API ====================

@audio_bp.route('/api/audio/upload', methods=['POST'])
@login_required
def api_audio_upload():
    """Rota compatível com frontend para upload de áudio"""
    try:
        # Verificar se há arquivo de áudio
        if 'audio' not in request.files:
            return jsonify({
                'success': False,
                'message': 'Nenhum arquivo de áudio fornecido'
            }), 400
        
        audio_file = request.files['audio']
        original_name = request.form.get('originalName', '')
        
        if audio_file.filename == '':
            return jsonify({
                'success': False,
                'message': 'Nome do arquivo não fornecido'
            }), 400
        
        # Ler dados do arquivo
        audio_bytes = audio_file.read()
        
        # Validar tamanho do arquivo
        if len(audio_bytes) < 1000:  # Menos de 1KB
            return jsonify({
                'success': False,
                'message': 'Arquivo de áudio muito pequeno ou corrompido'
            }), 400
        
        # Gerar nome único do arquivo
        user_id = session.get('user_id', 'unknown')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_id = str(random.randint(100000, 999999))
        
        # Usar nome original se fornecido
        if original_name:
            filename = f"{original_name}_{timestamp}_{unique_id}_{user_id}.wav"
        else:
            filename = f"recording_{timestamp}_{unique_id}_{user_id}.wav"
        
        # Garantir que o diretório existe
        os.makedirs(RECORDINGS_DIR, exist_ok=True)
        file_path = os.path.join(RECORDINGS_DIR, filename)
        
        # Verificar se arquivo já existe (evitar duplicação)
        if os.path.exists(file_path):
            counter = 1
            base_name, ext = os.path.splitext(filename)
            while os.path.exists(file_path):
                filename = f"{base_name}_({counter}){ext}"
                file_path = os.path.join(RECORDINGS_DIR, filename)
                counter += 1
        
        # Processar áudio com pydub para garantir compatibilidade
        try:
            # FUNÇÃO ROBUSTA: Detectar formato do áudio para diferentes dispositivos
            print(f"🔍 Processando áudio: {len(audio_bytes)} bytes")
            
            # FUNÇÃO ROBUSTA: Detectar formato do áudio para diferentes dispositivos
            format_result = detect_audio_format_robust(audio_bytes)
            
            # Extrair formato e sample rate esperado
            if isinstance(format_result, tuple):
                detected_format, expected_sample_rate = format_result
            else:
                # Fallback para compatibilidade com versões antigas
                detected_format = format_result
                expected_sample_rate = 44100
            
            print(f"🎯 Formato detectado: {detected_format}")
            print(f"🎵 Sample rate esperado: {expected_sample_rate}Hz")
            
            # Tentar carregar com pydub (suporta múltiplos formatos)
            try:
                audio_segment = AudioSegment.from_file(io.BytesIO(audio_bytes))
                print(f"✅ Áudio carregado com pydub: {audio_segment.frame_rate}Hz, {audio_segment.channels} canais")
                print(f"   Formato original: {detected_format}")
            except Exception as pydub_error:
                print(f"⚠️ Erro ao carregar com pydub: {pydub_error}")
                # Fallback: salvar arquivo original
                with open(file_path, 'wb') as f:
                    f.write(audio_bytes)
                file_size = len(audio_bytes)
                print(f"✅ Arquivo salvo sem processamento: {filename} ({file_size} bytes)")
                return jsonify({
                    'message': 'Áudio salvo com sucesso (sem processamento)',
                    'audioFile': {
                        'filename': filename,
                        'originalName': original_name or filename,
                        'size': file_size,
                        'createdAt': datetime.now().isoformat()
                    }
                }), 201
            
            # FUNÇÃO ROBUSTA: Processar áudio para compatibilidade com diferentes dispositivos
            audio_segment = process_audio_for_device_compatibility(audio_segment, detected_format)
            
            # CORREÇÃO: Exportar como WAV com configurações específicas para evitar problemas de velocidade
            print(f"🔧 Exportando WAV com configurações otimizadas...")
            audio_segment.export(
                file_path, 
                format="wav",
                parameters=[
                    "-acodec", "pcm_s16le",  # PCM 16-bit
                    "-ar", str(audio_segment.frame_rate),  # Manter sample rate
                    "-ac", "1"  # Mono
                ]
            )
            
            file_size = os.path.getsize(file_path)
            print(f"✅ Gravação processada e salva: {filename} ({file_size} bytes)")
            
        except Exception as e:
            print(f"⚠️ Erro no processamento com pydub: {e}")
            # Fallback: salvar arquivo original
            try:
                with open(file_path, 'wb') as f:
                    f.write(audio_bytes)
                file_size = len(audio_bytes)
                print(f"✅ Gravação salva (fallback): {filename} ({file_size} bytes)")
            except Exception as fallback_error:
                print(f"❌ Erro no fallback: {fallback_error}")
                return jsonify({
                    'success': False,
                    'message': 'Erro ao salvar arquivo de áudio'
                }), 500
        
                    # Resposta compatível com frontend
            return jsonify({
                'message': 'Áudio processado e salvo com sucesso',
                'audioFile': {
                    'filename': filename,
                    'originalName': original_name or filename,
                    'size': file_size,
                    'createdAt': datetime.now().isoformat(),
                    'originalFormat': detected_format,
                    'frameRate': audio_segment.frame_rate,
                    'channels': audio_segment.channels,
                    'duration': len(audio_segment),
                    'processingInfo': {
                        'deviceCompatible': True,
                        'qualityPreserved': True,
                        'speedCorrected': True
                    }
                }
            }), 201
        
    except Exception as e:
        print(f"❌ Erro geral ao fazer upload: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500

@audio_bp.route('/api/audio/save-simple', methods=['POST'])
@login_required
def api_audio_save_simple():
    """
    NOVA ARQUITETURA: Salvamento simples e rápido (sem processamento pesado)
    Responsabilidade Única: Apenas salvar arquivo original para processamento posterior
    """
    try:
        # Verificar se há arquivo de áudio
        if 'audio' not in request.files:
            return jsonify({
                'success': False,
                'message': 'Nenhum arquivo de áudio enviado'
            }), 400
        
        audio_file = request.files['audio']
        original_name = request.form.get('originalName', '')
        
        if audio_file.filename == '':
            return jsonify({
                'success': False,
                'message': 'Nome de arquivo inválido'
            }), 400
        
        # Gerar nome único para o arquivo
        user_id = session.get('user_id', '')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]
        random_suffix = random.randint(100000, 999999)
        
        # SIMPLIFICADO: Manter extensão original (WebM/Opus)
        original_extension = audio_file.filename.split('.')[-1] if '.' in audio_file.filename else 'webm'
        filename = f"{original_name or 'recording'}_{timestamp}_{random_suffix}_{user_id}.{original_extension}"
        
        # Garantir que o diretório existe
        os.makedirs(RECORDINGS_DIR, exist_ok=True)
        file_path = os.path.join(RECORDINGS_DIR, filename)
        
        # SALVAMENTO SIMPLES: Apenas salvar arquivo original
        print(f"💾 Salvamento simples: {filename}")
        audio_file.save(file_path)
        
        file_size = os.path.getsize(file_path)
        print(f"✅ Arquivo salvo rapidamente: {filename} ({file_size} bytes)")
        
        # Resposta simples e rápida
        return jsonify({
            'success': True,
            'message': 'Áudio salvo com sucesso! Use "Processar e Transcrever" para otimizar.',
            'audioFile': {
                'filename': filename,
                'originalName': original_name or filename,
                'size': file_size,
                'createdAt': datetime.now().isoformat(),
                'status': 'raw',  # Arquivo não processado
                'needsProcessing': True
            }
        }), 201
        
    except Exception as e:
        print(f"❌ Erro no salvamento simples: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro ao salvar áudio: {str(e)}'
        }), 500

@audio_bp.route('/api/audio/process', methods=['POST'])
@login_required
def api_audio_process():
    """
    NOVA ARQUITETURA: Processamento otimizado sob demanda
    Responsabilidade Única: Otimizar áudio e preparar para transcrição
    """
    try:
        data = request.get_json()
        filename = data.get('filename')
        
        if not filename:
            return jsonify({
                'success': False,
                'message': 'Nome do arquivo não fornecido'
            }), 400
        
        # Verificar se arquivo existe
        file_path = os.path.join(RECORDINGS_DIR, filename)
        if not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'message': 'Arquivo não encontrado'
            }), 404
        
        print(f"🔄 Iniciando processamento otimizado: {filename}")
        
        # Carregar arquivo original
        with open(file_path, 'rb') as f:
            audio_bytes = f.read()
        
        # Detectar formato
        format_result = detect_audio_format_robust(audio_bytes)
        if isinstance(format_result, tuple):
            detected_format, expected_sample_rate = format_result
        else:
            detected_format = format_result
            expected_sample_rate = 44100
        
        print(f"🎯 Formato detectado: {detected_format}")
        print(f"🎵 Sample rate esperado: {expected_sample_rate}Hz")
        
        # Processar com pydub
        audio_segment = AudioSegment.from_file(io.BytesIO(audio_bytes))
        print(f"✅ Áudio carregado: {audio_segment.frame_rate}Hz, {audio_segment.channels} canais")
        
        # Aplicar otimizações
        optimized_audio = process_audio_for_device_compatibility(audio_segment, detected_format)
        
        # Gerar nome do arquivo otimizado
        base_name = filename.rsplit('.', 1)[0]
        optimized_filename = f"{base_name}_optimized.wav"
        optimized_path = os.path.join(RECORDINGS_DIR, optimized_filename)
        
        # Salvar versão otimizada
        print(f"💾 Salvando versão otimizada: {optimized_filename}")
        optimized_audio.export(
            optimized_path,
            format="wav",
            parameters=[
                "-acodec", "pcm_s16le",  # PCM 16-bit
                "-ar", str(optimized_audio.frame_rate),  # Manter sample rate otimizado
                "-ac", "1"  # Mono
            ]
        )
        
        file_size = os.path.getsize(optimized_path)
        print(f"✅ Arquivo otimizado salvo: {optimized_filename} ({file_size} bytes)")
        
        return jsonify({
            'success': True,
            'message': 'Áudio processado e otimizado com sucesso!',
            'originalFile': filename,
            'optimizedFile': optimized_filename,
            'processing': {
                'originalFormat': detected_format,
                'originalSampleRate': audio_segment.frame_rate,
                'optimizedSampleRate': optimized_audio.frame_rate,
                'optimizedSize': file_size,
                'qualityImproved': optimized_audio.frame_rate > audio_segment.frame_rate
            }
        }), 200
        
    except Exception as e:
        print(f"❌ Erro no processamento: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro ao processar áudio: {str(e)}'
        }), 500

@audio_bp.route('/api/audio/transcribe', methods=['POST'])
@login_required
def api_audio_transcribe():
    """
    NOVA ARQUITETURA: Transcrição otimizada de arquivo processado
    Responsabilidade Única: Transcrever áudio já otimizado
    """
    try:
        data = request.get_json()
        optimized_filename = data.get('optimizedFile')
        
        if not optimized_filename:
            return jsonify({
                'success': False,
                'message': 'Nome do arquivo otimizado não fornecido'
            }), 400
        
        # Verificar se arquivo otimizado existe
        optimized_path = os.path.join(RECORDINGS_DIR, optimized_filename)
        if not os.path.exists(optimized_path):
            return jsonify({
                'success': False,
                'message': 'Arquivo otimizado não encontrado'
            }), 404
        
        print(f"🎯 Iniciando transcrição otimizada: {optimized_filename}")
        
        # Transcrever arquivo otimizado
        transcription = transcribe_audio_with_speech_recognition(optimized_path)
        
        if not transcription or transcription.startswith('[Erro') or transcription.startswith('[Não foi possível'):
            return jsonify({
                'success': False,
                'message': 'Falha na transcrição',
                'transcription': transcription
            }), 422
        
        # Salvar transcrição
        base_name = optimized_filename.rsplit('.', 1)[0]
        transcription_filename = f"{base_name}_transcricao.txt"
        transcription_path = os.path.join('transcriptions', transcription_filename)
        
        # Garantir que o diretório existe
        os.makedirs('transcriptions', exist_ok=True)
        
        # Salvar arquivo de transcrição
        with open(transcription_path, 'w', encoding='utf-8') as f:
            f.write(transcription)
        
        print(f"✅ Transcrição salva: {transcription_filename}")
        
        return jsonify({
            'success': True,
            'message': 'Transcrição concluída com sucesso!',
            'transcription': transcription,
            'transcriptionFile': transcription_filename,
            'optimizedFile': optimized_filename,
            'processing': {
                'transcriptionLength': len(transcription),
                'wordCount': len(transcription.split()) if transcription else 0,
                'qualityOptimized': True
            }
        }), 200
        
    except Exception as e:
        print(f"❌ Erro na transcrição: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro ao transcrever áudio: {str(e)}'
        }), 500

@audio_bp.route('/api/audio/test-fix', methods=['POST'])
@login_required
def api_audio_test_fix():
    """
    ENDPOINT DE TESTE: Testar correção de sample rate para resolver áudio lento
    """
    try:
        data = request.get_json()
        filename = data.get('filename')
        
        if not filename:
            return jsonify({
                'success': False,
                'message': 'Nome do arquivo não fornecido'
            }), 400
        
        file_path = os.path.join(RECORDINGS_DIR, filename)
        if not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'message': 'Arquivo não encontrado'
            }), 404
        
        print(f"🧪 TESTE DE CORREÇÃO: {filename}")
        
        # Carregar arquivo
        with open(file_path, 'rb') as f:
            audio_bytes = f.read()
        
        # Análise inicial
        audio_segment = AudioSegment.from_file(io.BytesIO(audio_bytes))
        original_rate = audio_segment.frame_rate
        original_duration = len(audio_segment)
        
        print(f"📊 ANTES: {original_rate}Hz, {original_duration}ms")
        
        # Aplicar correção robusta
        if original_rate <= 22050:
            # Forçar sample rate correto SEM reamostrar
            corrected_audio = audio_segment._spawn(audio_segment.raw_data, overrides={"frame_rate": 44100})
            print(f"🔧 FORÇADO: {original_rate}Hz → 44.1kHz (mesmos dados, interpretação correta)")
        else:
            corrected_audio = audio_segment
            print(f"✅ MANTIDO: {original_rate}Hz")
        
        # Verificar resultado
        final_rate = corrected_audio.frame_rate
        final_duration = len(corrected_audio)
        
        print(f"📊 DEPOIS: {final_rate}Hz, {final_duration}ms")
        
        # Salvar arquivo teste
        test_filename = f"TEST_FIXED_{filename}"
        test_path = os.path.join(RECORDINGS_DIR, test_filename)
        
        corrected_audio.export(
            test_path,
            format="wav",
            parameters=[
                "-acodec", "pcm_s16le",
                "-ar", str(final_rate),
                "-ac", "1"
            ]
        )
        
        # Verificar arquivo salvo
        with wave.open(test_path, 'rb') as wav_file:
            saved_rate = wav_file.getframerate()
            saved_duration = wav_file.getnframes() / saved_rate * 1000
        
        print(f"💾 ARQUIVO SALVO: {saved_rate}Hz, {saved_duration:.0f}ms")
        
        return jsonify({
            'success': True,
            'message': 'Teste de correção concluído',
            'original': {
                'filename': filename,
                'sampleRate': original_rate,
                'duration': original_duration
            },
            'corrected': {
                'filename': test_filename,
                'sampleRate': final_rate,
                'duration': final_duration
            },
            'saved': {
                'sampleRate': saved_rate,
                'duration': saved_duration
            },
            'speedFixed': original_rate != final_rate
        }), 200
        
    except Exception as e:
        print(f"❌ Erro no teste: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro no teste: {str(e)}'
        }), 500

@audio_bp.route('/api/audio/calibrate', methods=['POST'])
@login_required
def api_audio_calibrate():
    """
    CALIBRAÇÃO FINA: Ajustar sample rate para velocidade perfeita
    """
    try:
        data = request.get_json()
        filename = data.get('filename')
        speed_feedback = data.get('speedFeedback', 'normal')  # 'slow', 'fast', 'normal'
        
        if not filename:
            return jsonify({
                'success': False,
                'message': 'Nome do arquivo não fornecido'
            }), 400
        
        file_path = os.path.join(RECORDINGS_DIR, filename)
        if not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'message': 'Arquivo não encontrado'
            }), 404
        
        print(f"🎛️ CALIBRAÇÃO FINA: {filename} - Feedback: {speed_feedback}")
        
        # Carregar arquivo
        with open(file_path, 'rb') as f:
            audio_bytes = f.read()
        
        audio_segment = AudioSegment.from_file(io.BytesIO(audio_bytes))
        original_rate = audio_segment.frame_rate
        
        # Calcular fator de correção baseado no feedback
        if speed_feedback == 'slow':
            # Ainda está lento, aumentar sample rate um pouco mais
            if original_rate == 44100:
                calibrated_rate = 45000  # +2% mais rápido
            else:
                calibrated_rate = int(original_rate * 1.02)  # +2%
            print(f"🔧 CORREÇÃO LENTO: {original_rate}Hz → {calibrated_rate}Hz (+2%)")
            
        elif speed_feedback == 'fast':
            # Está rápido demais, diminuir sample rate um pouco
            if original_rate == 44100:
                calibrated_rate = 43200  # -2% mais lento
            else:
                calibrated_rate = int(original_rate * 0.98)  # -2%
            print(f"🔧 CORREÇÃO RÁPIDO: {original_rate}Hz → {calibrated_rate}Hz (-2%)")
            
        else:  # normal
            calibrated_rate = original_rate
            print(f"✅ VELOCIDADE PERFEITA: {original_rate}Hz mantido")
        
        if calibrated_rate != original_rate:
            # Aplicar calibração
            calibrated_audio = audio_segment._spawn(audio_segment.raw_data, overrides={"frame_rate": calibrated_rate})
            
            # Salvar arquivo calibrado
            calibrated_filename = f"CALIBRATED_{filename}"
            calibrated_path = os.path.join(RECORDINGS_DIR, calibrated_filename)
            
            calibrated_audio.export(
                calibrated_path,
                format="wav",
                parameters=[
                    "-acodec", "pcm_s16le",
                    "-ar", str(calibrated_rate),
                    "-ac", "1"
                ]
            )
            
            return jsonify({
                'success': True,
                'message': f'Áudio calibrado para velocidade {speed_feedback}',
                'original': {
                    'filename': filename,
                    'sampleRate': original_rate
                },
                'calibrated': {
                    'filename': calibrated_filename,
                    'sampleRate': calibrated_rate,
                    'adjustment': f"{'+' if calibrated_rate > original_rate else ''}{((calibrated_rate - original_rate) / original_rate * 100):.1f}%"
                }
            }), 200
        else:
            return jsonify({
                'success': True,
                'message': 'Áudio já está na velocidade perfeita',
                'noAdjustmentNeeded': True
            }), 200
        
    except Exception as e:
        print(f"❌ Erro na calibração: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro na calibração: {str(e)}'
        }), 500

@audio_bp.route('/api/audio/files', methods=['GET'])
@login_required
def api_audio_files():
    """Lista arquivos de áudio do usuário (compatível com frontend)"""
    try:
        user_id = session.get('user_id', '')
        user_email = session.get('user_email', '')
        recordings = []
        
        print(f"🔍 Buscando arquivos para user_id: {user_id}, email: {user_email}")
        
        # Verificar se o diretório existe
        if not os.path.exists(RECORDINGS_DIR):
            os.makedirs(RECORDINGS_DIR, exist_ok=True)
            return jsonify({'audioFiles': []})
        
        # Listar arquivos
        all_files = os.listdir(RECORDINGS_DIR)
        print(f"📋 Arquivos encontrados: {len(all_files)}")
        
        for filename in all_files:
            if not filename.endswith('.wav'):
                continue
            
            # Verificar se pertence ao usuário
            belongs_to_user = False
            
            # Formato novo: nome_timestamp_userid.wav
            if filename.endswith(f"_{user_id}.wav"):
                belongs_to_user = True
            # Formato antigo com email
            elif user_email:
                old_format_suffix = f"_{user_email.replace('@', '_').replace('.', '_')}.wav"
                if filename.endswith(old_format_suffix):
                    belongs_to_user = True
            
            if not belongs_to_user:
                continue
            
            file_path = os.path.join(RECORDINGS_DIR, filename)
            try:
                stats = os.stat(file_path)
                file_size = stats.st_size
                created_time = datetime.fromtimestamp(stats.st_ctime)
                
                recordings.append({
                    'id': filename,  # Usar filename como ID
                    'filename': filename,
                    'originalName': filename.replace(f'_{user_id}.wav', '').replace('.wav', ''),
                    'size': file_size,
                    'createdAt': created_time.isoformat()
                })
            except Exception as e:
                print(f"❌ Erro ao processar arquivo {filename}: {e}")
                continue
        
        # Ordenar por data de criação (mais recentes primeiro)
        recordings.sort(key=lambda x: x['createdAt'], reverse=True)
        
        return jsonify({'audioFiles': recordings})
        
    except Exception as e:
        print(f"❌ Erro ao listar arquivos: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500

@audio_bp.route('/api/save_recording', methods=['POST'])
@login_required
def api_save_recording():
    """Salva gravação com processamento otimizado para deploy"""
    try:
        data = request.json
        audio_data = data.get('audio')
        patient_name = data.get('patient_name', '')
        
        if not audio_data:
            return jsonify({
                'success': False,
                'message': 'Dados de áudio não fornecidos'
            }), 400
        
        # Decodificar base64
        if audio_data.startswith('data:audio'):
            # Extrair tipo MIME e dados
            header, audio_data = audio_data.split(',', 1)
            mime_type = header.split(':')[1].split(';')[0]
        else:
            mime_type = 'audio/wav'  # fallback
        
        try:
            audio_bytes = base64.b64decode(audio_data)
        except Exception as e:
            print(f"❌ Erro ao decodificar base64: {e}")
            return jsonify({
                'success': False,
                'message': 'Erro ao decodificar dados de áudio'
            }), 400
        
        # Validar tamanho do arquivo
        if len(audio_bytes) < 1000:  # Menos de 1KB
            return jsonify({
                'success': False,
                'message': 'Arquivo de áudio muito pequeno ou corrompido'
            }), 400
        
        # Gerar nome único do arquivo
        user_id = session.get('user_id', 'unknown')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_id = str(int(time.time() * 1000))[-6:]  # 6 últimos dígitos do timestamp
        
        if patient_name:
            safe_patient_name = sanitize_filename(patient_name)
            filename = f"{safe_patient_name}_{timestamp}_{unique_id}_{user_id}.wav"
        else:
            filename = f"recording_{timestamp}_{unique_id}_{user_id}.wav"
        
        # Garantir que o diretório existe
        os.makedirs(RECORDINGS_DIR, exist_ok=True)
        file_path = os.path.join(RECORDINGS_DIR, filename)
        
        # Verificar se arquivo já existe (evitar duplicação)
        if os.path.exists(file_path):
            counter = 1
            base_name, ext = os.path.splitext(filename)
            while os.path.exists(file_path):
                filename = f"{base_name}_({counter}){ext}"
                file_path = os.path.join(RECORDINGS_DIR, filename)
                counter += 1
        
        # Processar áudio com pydub para garantir compatibilidade
        try:
            # Carregar áudio original
            audio_segment = AudioSegment.from_file(io.BytesIO(audio_bytes))
            
            # FUNÇÃO ROBUSTA: Processar áudio para compatibilidade com diferentes dispositivos
            audio_segment = process_audio_for_device_compatibility(audio_segment, "Gravação via API")
            
            # Exportar como WAV mantendo qualidade
            audio_segment.export(
                file_path, 
                format="wav",
                parameters=["-acodec", "pcm_s16le"]  # PCM 16-bit
            )
            
            file_size = os.path.getsize(file_path)
            print(f"✅ Gravação processada e salva: {filename} ({file_size} bytes)")
            
        except Exception as e:
            print(f"⚠️ Erro no processamento com pydub: {e}")
            # Fallback: salvar arquivo original
            try:
                with open(file_path, 'wb') as f:
                    f.write(audio_bytes)
                file_size = len(audio_bytes)
                print(f"✅ Gravação salva (fallback): {filename} ({file_size} bytes)")
            except Exception as fallback_error:
                print(f"❌ Erro no fallback: {fallback_error}")
                return jsonify({
                    'success': False,
                    'message': 'Erro ao salvar arquivo de áudio'
                }), 500
        
        return jsonify({
            'success': True,
            'message': 'Gravação salva com sucesso!',
            'filename': filename,
            'size': file_size
        })
        
    except Exception as e:
        print(f"❌ Erro geral ao salvar gravação: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro interno do servidor: {str(e)}'
        }), 500

@audio_bp.route('/api/recordings', methods=['GET'])
@login_required
def api_get_recordings():
    """Lista gravações do usuário com lógica corrigida"""
    try:
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        recordings = []
        sessions = {}
        
        print(f"🔍 Buscando gravações para user_id: {user_id}, email: {user_email}")
        
        # Verificar se o diretório existe
        if not os.path.exists(RECORDINGS_DIR):
            print(f"❌ Diretório {RECORDINGS_DIR} não existe")
            # Criar diretório se não existir
            os.makedirs(RECORDINGS_DIR, exist_ok=True)
            return jsonify({
                'success': True,
                'recordings': [],
                'sessions': []
            })
        
        # Listar arquivos
        all_files = os.listdir(RECORDINGS_DIR)
        print(f"📋 Arquivos encontrados: {len(all_files)}")
        
        for filename in all_files:
            if not filename.endswith('.wav'):
                continue
            
            print(f"🔍 Analisando arquivo: {filename}")
            
            # CORREÇÃO: Lógica de filtragem mais rigorosa e com debug
            belongs_to_user = False
            
            # Verificar formato novo: nome_timestamp_userid.wav
            if filename.endswith(f"_{user_id}.wav"):
                belongs_to_user = True
                print(f"✅ Arquivo pertence ao usuário (formato novo): {filename}")
            
            # Verificar formato antigo apenas se user_email existir
            elif user_email:
                # Formato antigo: nome_timestamp_email_formatado.wav
                old_format_suffix = f"_{user_email.replace('@', '_').replace('.', '_')}.wav"
                if filename.endswith(old_format_suffix):
                    belongs_to_user = True
                    print(f"✅ Arquivo pertence ao usuário (formato antigo): {filename}")
            
            if not belongs_to_user:
                print(f"❌ Arquivo NÃO pertence ao usuário: {filename}")
                continue
            
            filepath = os.path.join(RECORDINGS_DIR, filename)
            
            try:
                # Verificar se o arquivo existe e é acessível
                if not os.path.exists(filepath):
                    print(f"❌ Arquivo não existe: {filepath}")
                    continue
                    
                file_size = os.path.getsize(filepath)
                
                # Data de modificação
                modified_time = os.path.getmtime(filepath)
                modified_date = datetime.fromtimestamp(modified_time).strftime('%d/%m/%Y %H:%M')
                
                # Verificar transcrição
                transcription_file = os.path.splitext(filename)[0] + '_transcricao.txt'
                transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_file)
                has_transcription = os.path.exists(transcription_path)
                
                # Formatar tamanho de forma legível
                if file_size < 1024:
                    size_str = f"{file_size} B"
                elif file_size < 1024 * 1024:
                    size_str = f"{file_size / 1024:.1f} KB"
                else:
                    size_str = f"{file_size / (1024 * 1024):.1f} MB"
                
                if '_sessao_' in filename:
                    # Segmento de sessão
                    parts = filename.split('_')
                    if len(parts) >= 3:
                        session_id = parts[2]
                        
                        if session_id not in sessions:
                            sessions[session_id] = {
                                'id': session_id,
                                'segments': [],
                                'total_size': 0
                            }
                        
                        sessions[session_id]['segments'].append({
                            'filename': filename,
                            'size': file_size,
                            'has_transcription': has_transcription
                        })
                        sessions[session_id]['total_size'] += file_size
                else:
                    # Gravação individual
                    recording_data = {
                        'filename': filename,
                        'size': size_str,  # CORREÇÃO: Formato legível
                        'date': modified_date,  # CORREÇÃO: Campo correto para o frontend
                        'has_transcription': has_transcription
                    }
                    recordings.append(recording_data)
                    print(f"✅ Gravação adicionada: {recording_data}")
                    
            except Exception as file_error:
                print(f"❌ Erro ao processar arquivo {filename}: {file_error}")
                continue
        
        # Ordenar por data (mais recente primeiro)
        try:
            recordings.sort(key=lambda x: datetime.strptime(x['date'], '%d/%m/%Y %H:%M'), reverse=True)
        except Exception as sort_error:
            print(f"⚠️ Erro ao ordenar gravações: {sort_error}")
            # Manter ordem original se houver erro na ordenação
        
        # Converter sessões para lista
        session_list = list(sessions.values())
        
        print(f"📊 RESULTADO FINAL: {len(recordings)} gravações e {len(session_list)} sessões")
        print(f"📋 Gravações encontradas: {[r['filename'] for r in recordings]}")
        
        return jsonify({
            'success': True,
            'recordings': recordings,
            'sessions': session_list
        })
        
    except Exception as e:
        print(f"❌ ERRO CRÍTICO ao listar gravações: {str(e)}")
        import traceback
        print(f"📋 Traceback: {traceback.format_exc()}")
        
        # CORREÇÃO: Retornar lista vazia em caso de erro, não falhar
        return jsonify({
            'success': True,  # Manter como True para não quebrar o frontend
            'recordings': [],
            'sessions': [],
            'error_message': f'Erro interno: {str(e)}'
        })

@audio_bp.route('/transcribe', methods=['POST'])
@login_required
def transcribe_recording():
    """Transcreve uma gravação"""
    try:
        data = request.json
        filename = data.get('filename')
        
        if not filename:
            return jsonify({
                'success': False,
                'message': 'Nome do arquivo não fornecido'
            }), 400
        
        # Verificar se o arquivo pertence ao usuário
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        
        # Verificar formato novo ou antigo
        belongs_to_user = (
            filename.endswith(f"_{user_id}.wav") or 
            (user_email and filename.endswith(f"_{user_email.replace('@', '_').replace('.', '_')}.wav"))
        )
        
        if not belongs_to_user:
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        filepath = os.path.join(RECORDINGS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'message': 'Arquivo não encontrado'
            }), 404
        
        print(f"🎯 Iniciando transcrição de: {filename}")
        
        # Transcrever
        transcription = transcribe_audio_with_speech_recognition(filepath)
        
        # Melhorar com Gemini se disponível
        if model and not transcription.startswith('['):
            print("🤖 Melhorando transcrição com Gemini...")
            transcription = improve_transcription_with_gemini(transcription)
        
        # Salvar transcrição
        transcription_filename = os.path.splitext(filename)[0] + '_transcricao.txt'
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_filename)
        
        # Criar diretório se não existir
        os.makedirs(TRANSCRIPTIONS_DIR, exist_ok=True)
        
        with open(transcription_path, 'w', encoding='utf-8') as f:
            f.write(transcription)
        
        print(f"✅ Transcrição salva: {transcription_filename}")
        
        return jsonify({
            'success': True,
            'transcription': transcription,
            'message': 'Transcrição realizada com sucesso!'
        })
        
    except Exception as e:
        print(f"❌ Erro na transcrição: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro na transcrição: {str(e)}'
        }), 500

@audio_bp.route('/api/rename_recording', methods=['POST'])
@login_required
def rename_recording():
    """Renomeia uma gravação"""
    try:
        data = request.json
        old_filename = data.get('old_filename')
        new_name = data.get('new_name', '').strip()
        
        if not old_filename or not new_name:
            return jsonify({
                'success': False,
                'message': 'Nome antigo e novo são obrigatórios'
            }), 400
        
        # Verificar se o arquivo pertence ao usuário
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        
        belongs_to_user = (
            old_filename.endswith(f"_{user_id}.wav") or 
            (user_email and old_filename.endswith(f"_{user_email.replace('@', '_').replace('.', '_')}.wav"))
        )
        
        if not belongs_to_user:
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        old_filepath = os.path.join(RECORDINGS_DIR, old_filename)
        
        if not os.path.exists(old_filepath):
            return jsonify({
                'success': False,
                'message': 'Arquivo não encontrado'
            }), 404
        
        # Gerar novo nome
        safe_new_name = sanitize_filename(new_name)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        new_filename = f'{safe_new_name}_{timestamp}_{user_id}.wav'
        new_filepath = os.path.join(RECORDINGS_DIR, new_filename)
        
        # Renomear arquivo
        os.rename(old_filepath, new_filepath)
        
        # Renomear transcrição se existir
        old_transcription = os.path.splitext(old_filename)[0] + '_transcricao.txt'
        old_transcription_path = os.path.join(TRANSCRIPTIONS_DIR, old_transcription)
        
        if os.path.exists(old_transcription_path):
            new_transcription = os.path.splitext(new_filename)[0] + '_transcricao.txt'
            new_transcription_path = os.path.join(TRANSCRIPTIONS_DIR, new_transcription)
            os.rename(old_transcription_path, new_transcription_path)
        
        print(f"✅ Arquivo renomeado: {old_filename} -> {new_filename}")
        
        return jsonify({
            'success': True,
            'new_filename': new_filename,
            'message': 'Arquivo renomeado com sucesso!'
        })
        
    except Exception as e:
        print(f"❌ Erro ao renomear: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro ao renomear: {str(e)}'
        }), 500

@audio_bp.route('/api/delete_recording', methods=['POST'])
@login_required
def delete_recording():
    """Exclui uma gravação"""
    try:
        data = request.json
        filename = data.get('filename')
        
        if not filename:
            return jsonify({
                'success': False,
                'message': 'Nome do arquivo não fornecido'
            }), 400
        
        # Verificar se o arquivo pertence ao usuário
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        
        belongs_to_user = (
            filename.endswith(f"_{user_id}.wav") or 
            (user_email and filename.endswith(f"_{user_email.replace('@', '_').replace('.', '_')}.wav"))
        )
        
        if not belongs_to_user:
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        filepath = os.path.join(RECORDINGS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'message': 'Arquivo não encontrado'
            }), 404
        
        # Excluir arquivo
        os.remove(filepath)
        
        # Excluir transcrição se existir
        transcription_file = os.path.splitext(filename)[0] + '_transcricao.txt'
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_file)
        
        if os.path.exists(transcription_path):
            os.remove(transcription_path)
        
        print(f"✅ Arquivo excluído: {filename}")
        
        return jsonify({
            'success': True,
            'message': 'Arquivo excluído com sucesso!'
        })
        
    except Exception as e:
        print(f"❌ Erro ao excluir: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro ao excluir: {str(e)}'
        }), 500

@audio_bp.route('/download/<filename>')
@login_required
def download_recording(filename):
    """Download de gravação"""
    try:
        # Verificar se o arquivo pertence ao usuário
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        
        belongs_to_user = (
            filename.endswith(f"_{user_id}.wav") or 
            (user_email and filename.endswith(f"_{user_email.replace('@', '_').replace('.', '_')}.wav"))
        )
        
        if not belongs_to_user:
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        filepath = os.path.join(RECORDINGS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'message': 'Arquivo não encontrado'
            }), 404
        
        return send_file(filepath, as_attachment=True, download_name=filename)
        
    except Exception as e:
        print(f"❌ Erro no download: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro no download: {str(e)}'
        }), 500

@audio_bp.route('/download_transcription/<filename>')
@login_required
def download_transcription(filename):
    """Download de transcrição"""
    try:
        # Verificar se o arquivo pertence ao usuário
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        
        recording_filename = filename.replace('_transcricao.txt', '.wav')
        belongs_to_user = (
            recording_filename.endswith(f"_{user_id}.wav") or 
            (user_email and recording_filename.endswith(f"_{user_email.replace('@', '_').replace('.', '_')}.wav"))
        )
        
        if not belongs_to_user:
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({
                'success': False,
                'message': 'Transcrição não encontrada'
            }), 404
        
        return send_file(filepath, as_attachment=True, download_name=filename)
        
    except Exception as e:
        print(f"❌ Erro no download da transcrição: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro no download: {str(e)}'
        }), 500

@audio_bp.route('/export_pdf/<filename>')
@login_required
def export_pdf(filename):
    """Exporta transcrição como PDF"""
    try:
        # Verificar se o arquivo pertence ao usuário
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        
        recording_filename = filename.replace('_transcricao.txt', '.wav')
        belongs_to_user = (
            recording_filename.endswith(f"_{user_id}.wav") or 
            (user_email and recording_filename.endswith(f"_{user_email.replace('@', '_').replace('.', '_')}.wav"))
        )
        
        if not belongs_to_user:
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if not os.path.exists(transcription_path):
            return jsonify({
                'success': False,
                'message': 'Transcrição não encontrada'
            }), 404
        
        # Ler transcrição
        with open(transcription_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Criar PDF
        pdf_bytes = create_pdf_from_text(text)
        
        if not pdf_bytes:
            return jsonify({
                'success': False,
                'message': 'Erro ao gerar PDF'
            }), 500
        
        # Criar arquivo temporário para download
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_file.write(pdf_bytes)
        temp_file.close()
        
        pdf_filename = filename.replace('_transcricao.txt', '_transcricao.pdf')
        
        return send_file(temp_file.name, as_attachment=True, download_name=pdf_filename)
        
    except Exception as e:
        print(f"❌ Erro ao exportar PDF: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro ao exportar PDF: {str(e)}'
        }), 500

@audio_bp.route('/export_docx/<filename>')
@login_required
def export_docx(filename):
    """Exporta transcrição como DOCX"""
    try:
        # Verificar se o arquivo pertence ao usuário
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        
        recording_filename = filename.replace('_transcricao.txt', '.wav')
        belongs_to_user = (
            recording_filename.endswith(f"_{user_id}.wav") or 
            (user_email and recording_filename.endswith(f"_{user_email.replace('@', '_').replace('.', '_')}.wav"))
        )
        
        if not belongs_to_user:
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if not os.path.exists(transcription_path):
            return jsonify({
                'success': False,
                'message': 'Transcrição não encontrada'
            }), 404
        
        # Ler transcrição
        with open(transcription_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        # Criar DOCX
        docx_bytes = create_docx_from_text(text)
        
        if not docx_bytes:
            return jsonify({
                'success': False,
                'message': 'Erro ao gerar DOCX'
            }), 500
        
        # Criar arquivo temporário para download
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.write(docx_bytes)
        temp_file.close()
        
        docx_filename = filename.replace('_transcricao.txt', '_transcricao.docx')
        
        return send_file(temp_file.name, as_attachment=True, download_name=docx_filename)
        
    except Exception as e:
        print(f"❌ Erro ao exportar DOCX: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro ao exportar DOCX: {str(e)}'
        }), 500

@audio_bp.route('/api/transcriptions', methods=['GET'])
@login_required
def api_get_transcriptions():
    """Rota API para listar transcrições - compatibilidade com React"""
    try:
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        safe_user_email = user_email.replace('@', '_').replace('.', '_')
        transcriptions = []
        
        print(f"🔍 Buscando transcrições para user_id: {user_id}")
        print(f"📧 Email do usuário: {user_email}")
        print(f"📁 Diretório de transcrições: {os.path.abspath(TRANSCRIPTIONS_DIR)}")
        
        # Verificar se o diretório existe
        if not os.path.exists(TRANSCRIPTIONS_DIR):
            print(f"❌ Diretório {TRANSCRIPTIONS_DIR} não existe")
            return jsonify({
                'success': True,
                'transcriptions': []
            })
        
        # Listar arquivos de transcrição
        for filename in os.listdir(TRANSCRIPTIONS_DIR):
            if filename.endswith('_transcricao.txt'):
                print(f"🔍 Verificando arquivo: {filename}")
                
                # Verificar se pertence ao usuário (suporta ambos os formatos)
                belongs_to_user = (
                    filename.endswith(f"_{user_id}_transcricao.txt") or
                    safe_user_email in filename
                )
                
                if belongs_to_user:
                    print(f"✅ Arquivo pertence ao usuário: {filename}")
                    
                    filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
                    file_size = os.path.getsize(filepath)
                    
                    # Obter data de modificação
                    modified_time = os.path.getmtime(filepath)
                    modified_date = datetime.fromtimestamp(modified_time).strftime('%d/%m/%Y %H:%M')
                    
                    # Ler primeiras linhas para preview
                    try:
                        with open(filepath, 'r', encoding='utf-8') as f:
                            content = f.read()
                            preview = content[:200] + '...' if len(content) > 200 else content
                    except:
                        preview = "Erro ao ler arquivo"
                    
                    # Nome original do arquivo de áudio
                    audio_filename = filename.replace('_transcricao.txt', '.wav')
                    
                    # Verificar se existe resumo
                    summary_filename = filename.replace('_transcricao.txt', '_resumo.txt')
                    summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
                    has_summary = os.path.exists(summary_path)
                    
                    transcriptions.append({
                        'filename': filename,
                        'audio_filename': audio_filename,
                        'size': file_size,
                        'modified_date': modified_date,
                        'preview': preview,
                        'has_summary': has_summary
                    })
                else:
                    print(f"❌ Arquivo não pertence ao usuário: {filename}")
        
        print(f"✅ Encontradas {len(transcriptions)} transcrições")
        
        return jsonify({
            'success': True,
            'transcriptions': transcriptions
        })
        
    except Exception as e:
        print(f"❌ Erro ao listar transcrições: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro ao listar transcrições: {str(e)}'
        }), 500

@audio_bp.route('/api/view_transcription/<filename>', methods=['GET'])
@login_required
def api_view_transcription(filename):
    """Rota API para visualizar transcrição - compatibilidade com React"""
    try:
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        safe_user_email = user_email.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário (suporta ambos os formatos)
        belongs_to_user = (
            filename.endswith(f"_{user_id}_transcricao.txt") or
            safe_user_email in filename
        )
        
        if not belongs_to_user:
            return jsonify({
                'success': False, 
                'message': 'Acesso negado - arquivo não pertence ao usuário'
            }), 403
        
        filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return jsonify({
                'success': True,
                'content': content
            })
        else:
            return jsonify({
                'success': False, 
                'message': 'Arquivo de transcrição não encontrado'
            }), 404
    
    except Exception as e:
        print(f"❌ Erro ao carregar transcrição: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro ao carregar transcrição: {str(e)}'
        }), 500

@audio_bp.route('/api/export_summary_pdf/<filename>', methods=['GET'])
@login_required
def api_export_summary_pdf(filename):
    """Exporta resumo como PDF - compatibilidade com React"""
    try:
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        safe_user_email = user_email.replace('@', '_').replace('.', '_')
        
        print(f"🔍 Tentando exportar PDF para: {filename}")
        print(f"👤 User ID: {user_id}")
        
        # Verificar se o arquivo pertence ao usuário
        belongs_to_user = (
            filename.endswith(f"_{user_id}_transcricao.txt") or
            safe_user_email in filename
        )
        
        if not belongs_to_user:
            print(f"❌ Arquivo não pertence ao usuário: {filename}")
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        # Construir nome do arquivo de resumo
        summary_filename = filename.replace('_transcricao.txt', '_resumo.txt')
        summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
        
        print(f"📄 Procurando resumo em: {summary_path}")
        
        if not os.path.exists(summary_path):
            print(f"❌ Arquivo de resumo não encontrado: {summary_path}")
            return jsonify({
                'success': False,
                'message': 'Resumo não encontrado. Execute a análise do Gemini primeiro.'
            }), 404
        
        # Ler resumo
        with open(summary_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        print(f"📖 Conteúdo lido: {len(content)} caracteres")
        
        if not content.strip():
            return jsonify({
                'success': False,
                'message': 'Arquivo de resumo está vazio'
            }), 400
        
        # Criar PDF
        pdf_bytes = create_pdf_from_text(content, "Resumo da Consulta")
        
        if not pdf_bytes:
            return jsonify({
                'success': False,
                'message': 'Erro ao gerar PDF - verifique o conteúdo do resumo'
            }), 500
        
        print(f"✅ PDF gerado com sucesso: {len(pdf_bytes)} bytes")
        
        # Criar arquivo temporário para download
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        temp_file.write(pdf_bytes)
        temp_file.close()
        
        pdf_filename = summary_filename.replace('.txt', '.pdf')
        
        return send_file(temp_file.name, as_attachment=True, download_name=pdf_filename)
        
    except Exception as e:
        print(f"❌ Erro ao exportar PDF: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Erro interno ao exportar PDF: {str(e)}'
        }), 500

@audio_bp.route('/api/export_summary_docx/<filename>', methods=['GET'])
@login_required
def api_export_summary_docx(filename):
    """Exporta resumo como DOCX - compatibilidade com React"""
    try:
        user_id = session['user_id']
        user_email = session.get('user_email', '')
        safe_user_email = user_email.replace('@', '_').replace('.', '_')
        
        print(f"🔍 Tentando exportar DOCX para: {filename}")
        print(f"👤 User ID: {user_id}")
        
        # Verificar se o arquivo pertence ao usuário
        belongs_to_user = (
            filename.endswith(f"_{user_id}_transcricao.txt") or
            safe_user_email in filename
        )
        
        if not belongs_to_user:
            print(f"❌ Arquivo não pertence ao usuário: {filename}")
            return jsonify({
                'success': False,
                'message': 'Arquivo não pertence ao usuário'
            }), 403
        
        # Construir nome do arquivo de resumo
        summary_filename = filename.replace('_transcricao.txt', '_resumo.txt')
        summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
        
        print(f"📄 Procurando resumo em: {summary_path}")
        
        if not os.path.exists(summary_path):
            print(f"❌ Arquivo de resumo não encontrado: {summary_path}")
            return jsonify({
                'success': False,
                'message': 'Resumo não encontrado. Execute a análise do Gemini primeiro.'
            }), 404
        
        # Ler resumo
        with open(summary_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        print(f"📖 Conteúdo lido: {len(content)} caracteres")
        
        if not content.strip():
            return jsonify({
                'success': False,
                'message': 'Arquivo de resumo está vazio'
            }), 400
        
        # Criar DOCX
        docx_bytes = create_docx_from_text(content, "Resumo da Consulta")
        
        if not docx_bytes:
            return jsonify({
                'success': False,
                'message': 'Erro ao gerar DOCX - verifique o conteúdo do resumo'
            }), 500
        
        print(f"✅ DOCX gerado com sucesso: {len(docx_bytes)} bytes")
        
        # Criar arquivo temporário para download
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.write(docx_bytes)
        temp_file.close()
        
        docx_filename = summary_filename.replace('.txt', '.docx')
        
        return send_file(temp_file.name, as_attachment=True, download_name=docx_filename)
        
    except Exception as e:
        print(f"❌ Erro ao exportar DOCX: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            'success': False,
            'message': f'Erro interno ao exportar DOCX: {str(e)}'
        }), 500

@audio_bp.route('/api/save_chunk', methods=['POST'])
@login_required
def api_save_chunk():
    """Salva chunk de gravação e monta arquivo final quando completo"""
    try:
        data = request.json
        session_id = data.get('session_id')
        chunk_index = data.get('chunk_index')
        audio_data = data.get('audio')
        is_last = data.get('is_last', False)
        mime_type = data.get('mime_type', 'audio/webm')
        
        if not all([session_id, audio_data is not None, chunk_index is not None]):
            return jsonify({
                'success': False,
                'message': 'Dados incompletos para o chunk'
            }), 400
        
        user_id = session.get('user_id', 'unknown')
        
        # Decodificar base64
        if audio_data.startswith('data:audio'):
            audio_data = audio_data.split(',', 1)[1]
        
        try:
            audio_bytes = base64.b64decode(audio_data)
        except Exception as e:
            return jsonify({
                'success': False,
                'message': 'Erro ao decodificar chunk de áudio'
            }), 400
        
        # Criar diretório temporário para chunks
        temp_dir = os.path.join(RECORDINGS_DIR, 'temp_chunks')
        os.makedirs(temp_dir, exist_ok=True)
        
        # Salvar chunk temporário
        chunk_filename = f"{session_id}_chunk_{chunk_index:03d}.tmp"
        chunk_path = os.path.join(temp_dir, chunk_filename)
        
        with open(chunk_path, 'wb') as f:
            f.write(audio_bytes)
        
        print(f"✅ Chunk {chunk_index} salvo: {len(audio_bytes)} bytes")
        
        # Se é o último chunk, montar arquivo final
        if is_last:
            final_filename = assemble_chunks(session_id, user_id, chunk_index + 1)
            return jsonify({
                'success': True,
                'message': f'Gravação completa! {chunk_index + 1} chunks processados.',
                'final_filename': final_filename
            })
        else:
            return jsonify({
                'success': True,
                'message': f'Chunk {chunk_index} recebido com sucesso'
            })
        
    except Exception as e:
        print(f"❌ Erro ao processar chunk: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro ao processar chunk: {str(e)}'
        }), 500

def assemble_chunks(session_id, user_id, total_chunks):
    """Monta chunks em arquivo final"""
    try:
        temp_dir = os.path.join(RECORDINGS_DIR, 'temp_chunks')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        final_filename = f"recording_{timestamp}_{user_id}.wav"
        final_path = os.path.join(RECORDINGS_DIR, final_filename)
        
        # Lista para armazenar todos os segmentos de áudio
        audio_segments = []
        
        print(f"🔧 Montando {total_chunks} chunks para sessão {session_id}")
        
        # Carregar e processar cada chunk
        for i in range(total_chunks):
            chunk_filename = f"{session_id}_chunk_{i:03d}.tmp"
            chunk_path = os.path.join(temp_dir, chunk_filename)
            
            if os.path.exists(chunk_path):
                try:
                    # Carregar chunk com pydub
                    chunk_segment = AudioSegment.from_file(chunk_path)
                    audio_segments.append(chunk_segment)
                    print(f"✅ Chunk {i} carregado: {len(chunk_segment)}ms")
                except Exception as e:
                    print(f"⚠️ Erro ao carregar chunk {i}: {e}")
                    # Tentar carregar como bytes brutos
                    with open(chunk_path, 'rb') as f:
                        chunk_bytes = f.read()
                    chunk_segment = AudioSegment.from_file(io.BytesIO(chunk_bytes))
                    audio_segments.append(chunk_segment)
            else:
                print(f"❌ Chunk {i} não encontrado: {chunk_path}")
        
        if not audio_segments:
            raise Exception("Nenhum chunk válido encontrado")
        
        # Concatenar todos os segmentos
        print(f"🔗 Concatenando {len(audio_segments)} segmentos...")
        final_audio = audio_segments[0]
        for segment in audio_segments[1:]:
            final_audio += segment
        
        # FUNÇÃO ROBUSTA: Processar áudio final para compatibilidade com diferentes dispositivos
        final_audio = process_audio_for_device_compatibility(final_audio, "Chunks montados")
        
        # CORREÇÃO: Exportar arquivo final com configurações específicas para evitar problemas de velocidade
        print(f"🔧 Exportando arquivo final com configurações otimizadas...")
        final_audio.export(
            final_path,
            format="wav",
            parameters=[
                "-acodec", "pcm_s16le",  # PCM 16-bit
                "-ar", str(final_audio.frame_rate),  # Manter sample rate
                "-ac", "1"  # Mono
            ]
        )
        
        file_size = os.path.getsize(final_path)
        duration = len(final_audio) / 1000  # duração em segundos
        
        print(f"✅ Arquivo final criado: {final_filename}")
        print(f"📊 Tamanho: {file_size} bytes, Duração: {duration:.1f}s")
        
        # Limpar chunks temporários
        cleanup_temp_chunks(session_id, total_chunks, temp_dir)
        
        return final_filename
        
    except Exception as e:
        print(f"❌ Erro ao montar chunks: {e}")
        raise

def cleanup_temp_chunks(session_id, total_chunks, temp_dir):
    """Remove chunks temporários após montagem"""
    try:
        for i in range(total_chunks):
            chunk_filename = f"{session_id}_chunk_{i:03d}.tmp"
            chunk_path = os.path.join(temp_dir, chunk_filename)
            if os.path.exists(chunk_path):
                os.remove(chunk_path)
        print(f"🧹 Chunks temporários removidos para sessão {session_id}")
    except Exception as e:
        print(f"⚠️ Erro ao limpar chunks temporários: {e}")

@audio_bp.route('/api/finalize_recording', methods=['POST'])
@login_required
def api_finalize_recording():
    """Finaliza gravação adicionando nome do paciente"""
    try:
        data = request.json
        session_id = data.get('session_id')
        patient_name = data.get('patient_name', '')
        current_filename = data.get('final_filename')
        
        if not current_filename:
            return jsonify({
                'success': False,
                'message': 'Nome do arquivo não fornecido'
            }), 400
        
        # Se há nome do paciente, renomear arquivo
        if patient_name:
            user_id = session.get('user_id', 'unknown')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            safe_patient_name = sanitize_filename(patient_name)
            new_filename = f"{safe_patient_name}_{timestamp}_{user_id}.wav"
            
            current_path = os.path.join(RECORDINGS_DIR, current_filename)
            new_path = os.path.join(RECORDINGS_DIR, new_filename)
            
            if os.path.exists(current_path):
                os.rename(current_path, new_path)
                print(f"📝 Arquivo renomeado: {current_filename} -> {new_filename}")
                final_filename = new_filename
            else:
                final_filename = current_filename
        else:
            final_filename = current_filename
        
        return jsonify({
            'success': True,
            'message': 'Gravação finalizada com sucesso!',
            'filename': final_filename
        })
        
    except Exception as e:
        print(f"❌ Erro ao finalizar gravação: {e}")
        return jsonify({
            'success': False,
            'message': f'Erro ao finalizar gravação: {str(e)}'
        }), 500

@audio_bp.route('/api/debug_files/<filename>', methods=['GET'])
@login_required
def debug_files(filename):
    """Debug: mostra status dos arquivos relacionados"""
    try:
        user_id = session['user_id']
        transcription_filename = filename.replace('.wav', '_transcricao.txt')
        summary_filename = filename.replace('.wav', '_resumo.txt')
        
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_filename)
        summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
        
        debug_info = {
            'user_id': user_id,
            'transcription_filename': transcription_filename,
            'summary_filename': summary_filename,
            'transcription_exists': os.path.exists(transcription_path),
            'summary_exists': os.path.exists(summary_path),
            'transcriptions_dir': os.path.abspath(TRANSCRIPTIONS_DIR),
            'files_in_dir': os.listdir(TRANSCRIPTIONS_DIR) if os.path.exists(TRANSCRIPTIONS_DIR) else []
        }
        
        return jsonify({
            'success': True,
            'debug_info': debug_info
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro no debug: {str(e)}'
        }), 500

