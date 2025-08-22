"""
Serviço de exportação de documentos
Separa as funcionalidades de exportação (PDF/DOCX) do app.py
"""

import io
from datetime import datetime
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ExportService:
    """Serviço para exportação de documentos em PDF e DOCX"""
    
    def __init__(self):
        # Registrar fonte que suporta acentos
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            self.font_name = 'DejaVuSans'
        except:
            self.font_name = 'Helvetica'
    
    def create_pdf_from_text(self, text, title="Resumo da Consulta"):
        """Cria um PDF a partir do texto"""
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=A4, 
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Estilos
            styles = getSampleStyleSheet()
            
            # Estilo personalizado para título
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=self.font_name,
                fontSize=16,
                spaceAfter=30,
                alignment=1,  # Centralizado
                textColor=colors.darkblue
            )
            
            # Estilo personalizado para corpo do texto
            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['Normal'],
                fontName=self.font_name,
                fontSize=11,
                spaceAfter=12,
                alignment=0,  # Justificado
                leftIndent=0,
                rightIndent=0
            )
            
            # Construir o documento
            story = []
            
            # Título
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 12))
            
            # Data
            data_atual = datetime.now().strftime("%d/%m/%Y às %H:%M")
            story.append(Paragraph(f"Gerado em: {data_atual}", body_style))
            story.append(Spacer(1, 20))
            
            # Processar o texto
            paragraphs = text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Escapar caracteres especiais para XML
                    safe_paragraph = paragraph.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe_paragraph, body_style))
                    story.append(Spacer(1, 12))
            
            # Construir PDF
            doc.build(story)
            
            buffer.seek(0)
            return buffer
        
        except Exception as e:
            print(f"Erro ao criar PDF: {e}")
            return None
    
    def create_docx_from_text(self, text, title="Resumo da Consulta"):
        """Cria um documento DOCX a partir do texto"""
        try:
            doc = Document()
            
            # Título
            title_paragraph = doc.add_heading(title, 0)
            title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Data
            data_atual = datetime.now().strftime("%d/%m/%Y às %H:%M")
            date_paragraph = doc.add_paragraph(f"Gerado em: {data_atual}")
            date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Espaço
            doc.add_paragraph()
            
            # Conteúdo
            paragraphs = text.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Salvar em buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
        
        except Exception as e:
            print(f"Erro ao criar DOCX: {e}")
            return None
    
    def export_text_as_pdf(self, text, filename, title="Resumo da Consulta"):
        """Exporta texto como PDF e retorna o buffer"""
        pdf_buffer = self.create_pdf_from_text(text, title)
        if pdf_buffer:
            return pdf_buffer
        else:
            raise Exception("Erro ao gerar PDF")
    
    def export_text_as_docx(self, text, filename, title="Resumo da Consulta"):
        """Exporta texto como DOCX e retorna o buffer"""
        docx_buffer = self.create_docx_from_text(text, title)
        if docx_buffer:
            return docx_buffer
        else:
            raise Exception("Erro ao gerar DOCX")
