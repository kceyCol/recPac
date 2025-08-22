from flask import Flask, render_template, request, jsonify, send_file, redirect, url_for, session, flash, send_from_directory
import os
from datetime import datetime
import base64
import wave
import io
import hashlib
import json
from functools import wraps
import re
from dotenv import load_dotenv
import google.generativeai as genai
import speech_recognition as sr
from pydub import AudioSegment
import tempfile
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
from flask_cors import CORS

# Importar blueprints
from auth import auth_bp, login_required
from audio_processing import audio_bp
from utils import configure_gemini, create_directories, RECORDINGS_DIR, TRANSCRIPTIONS_DIR

# Carregar variáveis de ambiente
load_dotenv()

# Configurar logging para produção
import logging
import sys
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    stream=sys.stdout
)
logger = logging.getLogger(__name__)
logger.info("🚀 Aplicação iniciando...")

app = Flask(__name__, static_folder=None, static_url_path=None)
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'sua_chave_secreta_aqui_mude_em_producao')

# Detectar se está em produção (Render)
is_production = os.getenv('RENDER') is not None

# Configurações de sessão adaptadas para produção
if is_production:
    app.config['SESSION_COOKIE_SECURE'] = False  # Temporariamente False para teste
    app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # Mudança para Lax
    allowed_origins = [
        "https://recpac.onrender.com",
        "http://localhost:3000"  # Para desenvolvimento local
    ]
else:
    app.config['SESSION_COOKIE_SECURE'] = False  # HTTP para desenvolvimento
    app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
    allowed_origins = ["http://localhost:3000"]

app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = 3600  # 1 hora

# Configurar CORS com origens dinâmicas
CORS(app, 
     supports_credentials=True,
     origins=allowed_origins,
     allow_headers=["Content-Type", "Authorization"],
     methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"]
)

# Registrar blueprints
app.register_blueprint(auth_bp)
app.register_blueprint(audio_bp)

# Criar diretórios necessários
create_directories()

# Configurar Gemini
gemini_api_key = os.getenv('GEMINI_API_KEY')
if gemini_api_key:
    genai.configure(api_key=gemini_api_key)
    model = genai.GenerativeModel('gemini-2.5-flash')
else:
    model = None
    print("⚠️ AVISO: GEMINI_API_KEY não configurada no arquivo .env")

USERS_FILE = 'users.json'

# Funções de autenticação (mantidas para compatibilidade)
def load_users():
    if os.path.exists(USERS_FILE):
        with open(USERS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_users(users):
    with open(USERS_FILE, 'w', encoding='utf-8') as f:
        json.dump(users, f, indent=2, ensure_ascii=False)

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\|?*]', '_', filename)

def load_default_prompt():
    """Carrega o prompt padrão do arquivo externo"""
    try:
        with open('default_prompt.txt', 'r', encoding='utf-8') as f:
            return f.read().strip()
    except FileNotFoundError:
        # Fallback caso o arquivo não exista
        return """Atue como um médico experiente, com vasta prática clínica, auxilie na avaliação de um paciente em consulta. A partir das informações clínicas que vou fornecer, extraia o máximo possível de detalhes relevantes para diagnóstico, conduta e planejamento terapêutico.

Organize as informações em tópicos claros, solicito uma análise detalhada, estruturada e prática, que possa servir de base para a tomada de decisão clínica."""

# Função para transcrever áudio - usando a implementação do audio_processing.py
def transcribe_audio_with_speech_recognition(audio_path):
    """Transcreve áudio usando a implementação robusta do audio_processing.py"""
    try:
        # Importar a função do módulo audio_processing
        from audio_processing import transcribe_audio_with_speech_recognition as transcribe_audio
        return transcribe_audio(audio_path)
    except ImportError:
        # Fallback se não conseguir importar
        return "[Erro: Módulo de transcrição não disponível]"
    except Exception as e:
        print(f"❌ Erro na transcrição: {e}")
        return f"[Erro na transcrição: {str(e)}]"

def improve_transcription_with_gemini(raw_transcription):
    """Melhora a transcrição usando Gemini"""
    try:
        if not model:
            return raw_transcription
        
        prompt = f"""
Você é um assistente especializado em melhorar transcrições de áudio médico.

Transcrição original:
{raw_transcription}

Por favor, corrija:
1. Erros de ortografia
2. Pontuação adequada
3. Formatação de parágrafos
4. Termos médicos que possam ter sido mal transcritos

Mantenha o conteúdo original, apenas melhore a qualidade da transcrição.
Responda apenas com a transcrição melhorada, sem comentários adicionais.
"""
        
        response = model.generate_content(prompt)
        improved_text = response.text.strip()
        
        if improved_text and len(improved_text) > 10:
            return improved_text
        else:
            return raw_transcription
    
    except Exception as e:
        print(f"Erro ao melhorar transcrição com Gemini: {e}")
        return raw_transcription

def create_pdf_from_text(text, title="Resumo da Consulta"):
    """Cria um PDF a partir do texto"""
    try:
        # Registrar fonte que suporta acentos
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            font_name = 'DejaVuSans'
        except:
            font_name = 'Helvetica'
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, 
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Estilos
        styles = getSampleStyleSheet()
        
        # Estilo personalizado para título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=font_name,
            fontSize=16,
            spaceAfter=30,
            alignment=1,  # Centralizado
            textColor=colors.darkblue
        )
        
        # Estilo personalizado para corpo do texto
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=font_name,
            fontSize=11,
            spaceAfter=12,
            alignment=0,  # Justificado
            leftIndent=0,
            rightIndent=0
        )
        
        # Construir o documento
        story = []
        
        # Título
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Data
        data_atual = datetime.now().strftime("%d/%m/%Y às %H:%M")
        story.append(Paragraph(f"Gerado em: {data_atual}", body_style))
        story.append(Spacer(1, 20))
        
        # Processar o texto
        paragraphs = text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                # Escapar caracteres especiais para XML
                safe_paragraph = paragraph.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(safe_paragraph, body_style))
                story.append(Spacer(1, 12))
        
        # Construir PDF
        doc.build(story)
        
        buffer.seek(0)
        return buffer
    
    except Exception as e:
        print(f"Erro ao criar PDF: {e}")
        return None

def create_docx_from_text(text, title="Resumo da Consulta"):
    """Cria um documento DOCX a partir do texto"""
    try:
        doc = Document()
        
        # Título
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data
        data_atual = datetime.now().strftime("%d/%m/%Y às %H:%M")
        date_paragraph = doc.add_paragraph(f"Gerado em: {data_atual}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Espaço
        doc.add_paragraph()
        
        # Conteúdo
        paragraphs = text.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Salvar em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    except Exception as e:
        print(f"Erro ao criar DOCX: {e}")
        return None

# Servir arquivos estáticos do React
@app.route('/static/<path:filepath>')
def serve_react_static(filepath):
    react_static = os.path.join('frontend', 'build', 'static')
    if os.path.exists(react_static):
        return send_from_directory(react_static, filepath)
    else:
        return jsonify({'error': 'Static file not found', 'file': filepath}), 404

# Rotas principais (mantidas para compatibilidade)
@app.route('/app')
@login_required
def app_route():
    return render_template('index.html', user_name=session.get('user_name', session.get('user_id')))

@app.route('/transcribe', methods=['POST'])
@login_required
def transcribe_recording():
    try:
        data = request.json
        filename = data['filename']
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(RECORDINGS_DIR, filename)
        
        if not os.path.exists(filepath):
            return jsonify({'success': False, 'message': 'Arquivo não encontrado'}), 404
        
        print(f"🎯 Iniciando transcrição de: {filename}")
        
        # Transcrever o áudio
        transcription = transcribe_audio_with_speech_recognition(filepath)
        
        # Melhorar com Gemini se disponível
        if model and not transcription.startswith('['):
            print("🤖 Melhorando transcrição com Gemini...")
            transcription = improve_transcription_with_gemini(transcription)
        
        # Salvar transcrição
        base_filename = os.path.splitext(filename)[0]
        transcription_filename = f'{base_filename}_transcricao.txt'
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_filename)
        
        with open(transcription_path, 'w', encoding='utf-8') as f:
            f.write(transcription)
        
        print(f"✅ Transcrição salva: {transcription_filename}")
        
        return jsonify({
            'success': True,
            'transcription': transcription,
            'transcription_file': transcription_filename
        })
    
    except Exception as e:
        print(f"❌ Erro na transcrição: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Erro na transcrição: {str(e)}'
        }), 500

@app.route('/download_transcription/<filename>')
@login_required
def download_transcription(filename):
    try:
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        return send_from_directory(TRANSCRIPTIONS_DIR, filename, as_attachment=True)
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 404

@app.route('/rename_recording', methods=['POST'])
@login_required
def rename_recording():
    try:
        data = request.json
        old_filename = data['old_filename']
        new_name = data['new_name']
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in old_filename or user_id in old_filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        # Sanitizar novo nome
        safe_new_name = sanitize_filename(new_name)
        if not safe_new_name:
            return jsonify({'success': False, 'message': 'Nome inválido'}), 400
        
        # Gerar novo nome do arquivo
        # Extrair timestamp e user_id do filename antigo
        parts = old_filename.replace('.wav', '').split('_')
        
        # Determinar formato do arquivo baseado na estrutura
        if len(parts) >= 3:
            # Formato: nome_timestamp_userid.wav
            timestamp = parts[-2]
            user_id_from_file = parts[-1]
            new_filename = f'{safe_new_name}_{timestamp}_{user_id_from_file}.wav'
        else:
            # Formato antigo ou simples
            timestamp = parts[-1] if len(parts) > 1 else datetime.now().strftime('%Y%m%d_%H%M%S')
            new_filename = f'{safe_new_name}_{timestamp}_{safe_user_id}.wav'
        
        old_path = os.path.join(RECORDINGS_DIR, old_filename)
        new_path = os.path.join(RECORDINGS_DIR, new_filename)
        
        if os.path.exists(old_path):
            os.rename(old_path, new_path)
            
            # Renomear transcrição se existir
            old_transcription = os.path.splitext(old_filename)[0] + '_transcricao.txt'
            new_transcription = os.path.splitext(new_filename)[0] + '_transcricao.txt'
            
            old_transcription_path = os.path.join(TRANSCRIPTIONS_DIR, old_transcription)
            new_transcription_path = os.path.join(TRANSCRIPTIONS_DIR, new_transcription)
            
            if os.path.exists(old_transcription_path):
                os.rename(old_transcription_path, new_transcription_path)
            
            return jsonify({
                'success': True,
                'message': 'Arquivo renomeado com sucesso!',
                'new_filename': new_filename
            })
        else:
            return jsonify({'success': False, 'message': 'Arquivo não encontrado'}), 404
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao renomear: {str(e)}'
        }), 500

@app.route('/transcriptions')
@login_required
def transcriptions_page():
    return render_template('transcriptions.html', user_name=session.get('user_name'))

@app.route('/api/transcriptions')
@login_required
def get_all_transcriptions():
    try:
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        transcriptions = []
        
        # Listar arquivos de transcrição
        for filename in os.listdir(TRANSCRIPTIONS_DIR):
            if filename.endswith('_transcricao.txt') and (user_id in filename or session.get('user_email', '').replace('@', '_').replace('.', '_') in filename):
                filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
                file_size = os.path.getsize(filepath)
                
                # Obter data de modificação
                modified_time = os.path.getmtime(filepath)
                modified_date = datetime.fromtimestamp(modified_time).strftime('%d/%m/%Y %H:%M')
                
                # Ler primeiras linhas para preview
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                        preview = content[:200] + '...' if len(content) > 200 else content
                except:
                    preview = "Erro ao ler arquivo"
                
                # Nome original do arquivo de áudio
                audio_filename = filename.replace('_transcricao.txt', '.wav')
                
                # Verificar se existe resumo
                summary_filename = filename.replace('_transcricao.txt', '_resumo.txt')
                summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
                has_summary = os.path.exists(summary_path)
                
                transcriptions.append({
                    'filename': filename,
                    'audio_filename': audio_filename,
                    'size': file_size,
                    'modified_date': modified_date,
                    'preview': preview,
                    'has_summary': has_summary
                })
        
        # Ordenar por data de modificação (mais recente primeiro)
        transcriptions.sort(key=lambda x: x['modified_date'], reverse=True)
        
        return jsonify({
            'success': True,
            'transcriptions': transcriptions
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao carregar transcrições: {str(e)}'
        }), 500

#---------------------------------------------------------------------------------
@app.route('/api/generate_summary', methods=['POST'])
@login_required
def generate_summary():
    try:
        data = request.get_json()
        transcription_content = data.get('transcription', '')
        custom_prompt = data.get('custom_prompt', '')
        filename = data.get('filename', '')  # Adicionar filename
        
        if not transcription_content:
            return jsonify({'error': 'Conteúdo da transcrição é obrigatório'}), 400
        
        # Verificar se o Gemini está configurado
        if not model:
            return jsonify({'error': 'Gemini AI não está configurado. Verifique a GEMINI_API_KEY no arquivo .env'}), 500
       
        try:
            # Usar prompt personalizado se fornecido, senão usar o padrão do arquivo
            if custom_prompt.strip():
                prompt = f"{transcription_content}\n\n{custom_prompt}"
            else:
                # Carregar prompt padrão do arquivo externo
                default_prompt = load_default_prompt()
                prompt = f"{transcription_content}\n\n{default_prompt}"
            
            # Gerar resumo com Gemini
            response = model.generate_content(prompt)
            summary = response.text
            
            # NOVO: Salvar o resumo automaticamente
            if filename:
                user_id = session['user_id']
                # Construir nome do arquivo de resumo
                if filename.endswith('.wav'):
                    summary_filename = filename.replace('.wav', f'_{user_id}_resumo.txt')
                elif filename.endswith('_transcricao.txt'):
                    summary_filename = filename.replace('_transcricao.txt', '_resumo.txt')
                else:
                    summary_filename = f"{filename}_{user_id}_resumo.txt"
                
                summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
                
                # Salvar arquivo de resumo
                with open(summary_path, 'w', encoding='utf-8') as f:
                    f.write(summary)
                
                print(f"✅ Resumo salvo em: {summary_path}")
            
            return jsonify({
                'summary': summary
            })
            
        except Exception as gemini_error:
            print(f"Erro do Gemini: {str(gemini_error)}")
            return jsonify({'error': f'Erro ao gerar resumo com Gemini: {str(gemini_error)}'}), 500
        
    except Exception as e:
        print(f"Erro geral: {str(e)}")
        return jsonify({'error': f'Erro interno: {str(e)}'}), 500

#---------------------------------------------------------------
@app.route('/api/view_summary/<filename>')
@login_required
def view_summary(filename):
    try:
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        # Construir nome do arquivo de resumo
        if filename.endswith('_transcricao.txt'):
            summary_filename = filename.replace('_transcricao.txt', '_resumo.txt')
        else:
            summary_filename = filename
        
        summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
        
        if os.path.exists(summary_path):
            with open(summary_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return jsonify({
                'success': True,
                'content': content,
                'filename': summary_filename
            })
        else:
            return jsonify({'success': False, 'message': 'Resumo não encontrado'}), 404
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao carregar resumo: {str(e)}'
        }), 500

@app.route('/export_summary_pdf/<filename>')
@login_required
def export_summary_pdf(filename):
    try:
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        # Construir nome do arquivo de resumo
        if filename.endswith('_transcricao.txt'):
            summary_filename = filename.replace('_transcricao.txt', '_resumo.txt')
        else:
            summary_filename = filename
        
        summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
        
        if not os.path.exists(summary_path):
            return jsonify({'success': False, 'message': 'Resumo não encontrado'}), 404
        
        with open(summary_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Criar PDF
        pdf_buffer = create_pdf_from_text(content, "Resumo da Consulta")
        
        if pdf_buffer:
            pdf_filename = summary_filename.replace('.txt', '.pdf')
            return send_file(
                pdf_buffer,
                as_attachment=True,
                download_name=pdf_filename,
                mimetype='application/pdf'
            )
        else:
            return jsonify({'success': False, 'message': 'Erro ao gerar PDF'}), 500
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao exportar PDF: {str(e)}'
        }), 500

@app.route('/export_summary_docx/<filename>')
@login_required
def export_summary_docx(filename):
    try:
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        # Construir nome do arquivo de resumo
        if filename.endswith('_transcricao.txt'):
            summary_filename = filename.replace('_transcricao.txt', '_resumo.txt')
        else:
            summary_filename = filename
        
        summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_filename)
        
        if not os.path.exists(summary_path):
            return jsonify({'success': False, 'message': 'Resumo não encontrado'}), 404
        
        with open(summary_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Criar DOCX
        docx_buffer = create_docx_from_text(content, "Resumo da Consulta")
        
        if docx_buffer:
            docx_filename = summary_filename.replace('.txt', '.docx')
            return send_file(
                docx_buffer,
                as_attachment=True,
                download_name=docx_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            return jsonify({'success': False, 'message': 'Erro ao gerar DOCX'}), 500
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao exportar DOCX: {str(e)}'
        }), 500

@app.route('/recordings')
@login_required
def get_recordings():
    try:
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        recordings = []
        sessions = {}
        
        # Listar arquivos de gravação
        for filename in os.listdir(RECORDINGS_DIR):
            if filename.endswith('.wav') and (safe_user_id in filename or user_id in filename):
                filepath = os.path.join(RECORDINGS_DIR, filename)
                file_size = os.path.getsize(filepath)
                
                # Verificar se existe transcrição
                transcription_file = os.path.splitext(filename)[0] + '_transcricao.txt'
                transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_file)
                has_transcription = os.path.exists(transcription_path)
                
                if '_sessao_' in filename:
                    # É um segmento de sessão
                    parts = filename.split('_')
                    session_id = parts[2]  # sessao_ID
                    
                    if session_id not in sessions:
                        sessions[session_id] = {
                            'id': session_id,
                            'segments': [],
                            'total_size': 0
                        }
                    
                    sessions[session_id]['segments'].append({
                        'filename': filename,
                        'size': file_size,
                        'has_transcription': has_transcription
                    })
                    sessions[session_id]['total_size'] += file_size
                else:
                    # Gravação individual
                    recordings.append({
                        'filename': filename,
                        'size': file_size,
                        'has_transcription': has_transcription
                    })
        
        # Converter sessões para lista
        sessions_list = list(sessions.values())
        
        return render_template('index.html', 
                             recordings=recordings, 
                             sessions=sessions_list,
                             user_name=session.get('user_name'))
    
    except Exception as e:
        flash(f'Erro ao carregar gravações: {str(e)}', 'error')
        return render_template('index.html', recordings=[], sessions=[])

@app.route('/delete_recording', methods=['POST'])
@login_required
def delete_recording():
    try:
        data = request.json
        filename = data['filename']
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(RECORDINGS_DIR, filename)
        
        if os.path.exists(filepath):
            os.remove(filepath)
            
            # Remover transcrição se existir
            transcription_file = os.path.splitext(filename)[0] + '_transcricao.txt'
            transcription_path = os.path.join(TRANSCRIPTIONS_DIR, transcription_file)
            
            if os.path.exists(transcription_path):
                os.remove(transcription_path)
            
            # Remover resumo se existir
            summary_file = os.path.splitext(filename)[0] + '_resumo.txt'
            summary_path = os.path.join(TRANSCRIPTIONS_DIR, summary_file)
            
            if os.path.exists(summary_path):
                os.remove(summary_path)
            
            return jsonify({
                'success': True,
                'message': 'Arquivo deletado com sucesso!'
            })
        else:
            return jsonify({'success': False, 'message': 'Arquivo não encontrado'}), 404
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao deletar: {str(e)}'
        }), 500

@app.route('/download/<filename>')
@login_required
def download_file(filename):
    try:
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        return send_from_directory(RECORDINGS_DIR, filename, as_attachment=True)
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 404

@app.route('/finalize_session', methods=['POST'])
@login_required
def finalize_session():
    try:
        data = request.json
        session_id = data['session_id']
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Buscar todos os segmentos da sessão
        segments = []
        for filename in os.listdir(RECORDINGS_DIR):
            if f'_sessao_{session_id}_' in filename and (safe_user_id in filename or user_id in filename):
                segments.append(filename)
        
        if not segments:
            return jsonify({'success': False, 'message': 'Nenhum segmento encontrado para esta sessão'}), 404
        
        # Ordenar segmentos por timestamp
        segments.sort()
        
        # Combinar áudios usando pydub
        combined_audio = AudioSegment.empty()
        
        for segment_filename in segments:
            segment_path = os.path.join(RECORDINGS_DIR, segment_filename)
            segment_audio = AudioSegment.from_wav(segment_path)
            combined_audio += segment_audio
        
        # Salvar áudio combinado
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        final_filename = f'Sessao_{session_id}_{timestamp}_{safe_user_id}.wav'
        final_path = os.path.join(RECORDINGS_DIR, final_filename)
        
        # CORREÇÃO: Exportar com configurações específicas para evitar problemas de velocidade
        combined_audio.export(final_path, format="wav", parameters=[
            "-acodec", "pcm_s16le",  # PCM 16-bit
            "-ar", "44100",  # Sample rate padrão
            "-ac", "1"  # Mono
        ])
        
        return jsonify({
            'success': True,
            'message': 'Sessão finalizada com sucesso!',
            'filename': final_filename
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao finalizar sessão: {str(e)}'
        }), 500

@app.route('/api/save_summary_copy', methods=['POST'])
@login_required
def save_summary_copy():
    try:
        data = request.json
        summary_content = data['summary_content']
        original_filename = data['original_filename']
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in original_filename or user_id in original_filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        # Gerar nome do arquivo de cópia
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_name = original_filename.replace('_resumo.txt', '').replace('_transcricao.txt', '')
        copy_filename = f'{base_name}_resumo_editado_{timestamp}.txt'
        copy_path = os.path.join(TRANSCRIPTIONS_DIR, copy_filename)
        
        # Salvar cópia editada
        with open(copy_path, 'w', encoding='utf-8') as f:
            f.write(summary_content)
        
        return jsonify({
            'success': True,
            'message': 'Cópia do resumo salva com sucesso!',
            'filename': copy_filename
        })
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao salvar cópia: {str(e)}'
        }), 500

@app.route('/view_transcription/<filename>')
@login_required
def view_transcription(filename):
    try:
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        filepath = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        if os.path.exists(filepath):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            
            return jsonify({
                'success': True,
                'content': content
            })
        else:
            return jsonify({'success': False, 'message': 'Arquivo de transcrição não encontrado'}), 404
    
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao carregar transcrição: {str(e)}'
        }), 500

@app.route('/api/save_transcription', methods=['POST'])
@login_required
def save_transcription():
    try:
        data = request.json
        filename = data['filename']
        content = data['content']
        user_id = session['user_id']
        safe_user_id = user_id.replace('@', '_').replace('.', '_')
        
        # Verificar se o arquivo pertence ao usuário
        if not (safe_user_id in filename or user_id in filename):
            return jsonify({'success': False, 'message': 'Acesso negado'}), 403
        
        # Salvar transcrição
        transcription_path = os.path.join(TRANSCRIPTIONS_DIR, filename)
        
        with open(transcription_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        return jsonify({
            'success': True,
            'message': 'Transcrição salva com sucesso!'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'Erro ao salvar transcrição: {str(e)}'
        }), 500

@app.route('/api/default_prompt')
@login_required
def get_default_prompt():
    """Retorna o prompt padrão para o frontend"""
    try:
        default_prompt = load_default_prompt()
        return jsonify({'prompt': default_prompt})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
# Rota para a página inicial
@app.route('/')
def index():
    return send_from_directory('frontend/build', 'index.html')

# Rota para arquivos estáticos do React
@app.route('/<path:path>')
def static_files(path):
    # Se for API, não interceptar
    if path.startswith('api'):
        from flask import abort
        abort(404)
    
    # Tentar servir o arquivo
    try:
        return send_from_directory('frontend/build', path)
    except:
        # Se não encontrar, servir index.html para SPA routing
        return send_from_directory('frontend/build', 'index.html')
        
if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    debug_mode = os.environ.get('FLASK_ENV') == 'development'
    app.run(host='0.0.0.0', port=port, debug=debug_mode)